import os
import uuid
import json
import re
from flask import Flask, render_template, request, send_file, flash, redirect, session, jsonify, url_for
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
import google.generativeai as genai
from playwright.sync_api import sync_playwright
import pdfplumber
import logging
from datetime import timedelta
import traceback
import urllib.parse
import bleach
from flask_session import Session

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Gemini API Key
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    raise RuntimeError("Missing GEMINI_API_KEY environment variable")

genai.configure(api_key=api_key)

app = Flask(__name__)
app.secret_key = 'supersecret'
app.permanent_session_lifetime = timedelta(minutes=30)
UPLOAD_FOLDER = 'Uploads'
GENERATED_FOLDER = 'generated'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Configure server-side sessions
app.config['SESSION_TYPE'] = 'filesystem'
Session(app)

# Register Jinja2 filter for date formatting
def jinja_format_date(date_str):
    return format_date_for_display(date_str)

app.jinja_env.filters['format_date'] = jinja_format_date

model = genai.GenerativeModel("gemini-1.5-flash")

VALID_SECTION_IDS = [
    'education-section', 'experience-section', 'summary-section', 'projects-section',
    'roles-section', 'skills-section', 'personal-details-section', 'work-experience-section'
]

# Allowed HTML tags and attributes for rich text fields
ALLOWED_TAGS = ['b', 'i', 'ul', 'ol', 'li']
ALLOWED_ATTRIBUTES = {}

# Excel styling
HEADER_FONT = Font(name='Arial', size=12, bold=True, color='FFFFFF')
HEADER_FILL = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
DATA_FONT = Font(name='Arial', size=11)
BORDER = Border(
    left=Side(style='thin'), 
    right=Side(style='thin'), 
    top=Side(style='thin'), 
    bottom=Side(style='thin')
)
CENTER_ALIGN = Alignment(horizontal='center', vertical='center')
LEFT_ALIGN = Alignment(horizontal='left', vertical='center', wrap_text=True)

def apply_cell_style(cell, is_header=False):
    if is_header:
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = CENTER_ALIGN
    else:
        cell.font = DATA_FONT
        cell.alignment = LEFT_ALIGN
    cell.border = BORDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(filepath, ext):
    try:
        if ext == 'pdf':
            with pdfplumber.open(filepath) as pdf:
                text = ''
                for page in pdf.pages:
                    text += page.extract_text() or ''
                    for table in page.extract_tables():
                        for row in table:
                            text += ' | '.join([cell or '' for cell in row]) + '\n'
                logging.info(f"Extracted text from PDF file '{filepath}': {text[:100]}...")
                return text
        elif ext == 'docx':
            doc = Document(filepath)
            text = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    text.append(f"# {para.text}")
                else:
                    text.append(para.text)
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    table_text.append(row_text)
                text.append('\n'.join(table_text))
            full_text = '\n'.join(line for line in text if line.strip())
            logging.info(f"Extracted text from DOCX file '{filepath}': {full_text[:100]}...")
            return full_text
        elif ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
                logging.info(f"Extracted text from TXT file '{filepath}': {text[:100]}...")
                return text
        return ""
    except Exception as e:
        logging.error(f"Error extracting text from {filepath}: {e}")
        return ""

def clean_formatting(text):
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'•', '-', text)
    text = re.sub(r'\n\s*-\s*', '\n- ', text)
    text = re.sub(r'^#+\s*(.*?)\s*$', r'# \1', text, flags=re.MULTILINE)
    
    # Date standardization
    month_map = {
        'jan': '01', 'january': '01', 'feb': '02', 'february': '02', 'mar': '03', 'march': '03',
        'apr': '04', 'april': '04', 'may': '05', 'jun': '06', 'june': '06', 'jul': '07', 'july': '07',
        'aug': '08', 'august': '08', 'sep': '09', 'sept': '09', 'september': '09', 
        'oct': '10', 'october': '10', 'nov': '11', 'november': '11', 'dec': '12', 'december': '12'
    }
    
    def standardize_date(match):
        date_str = match.group(0)
        if re.match(r'(\d{1,2})[/-](\d{4})', date_str):
            month, year = re.match(r'(\d{1,2})[/-](\d{4})', date_str).groups()
            month = month.zfill(2)
            return f"{year}-{month}"
        elif re.match(r'(\d{4})[/-](\d{1,2})', date_str):
            year, month = re.match(r'(\d{4})[/-](\d{1,2})', date_str).groups()
            month = month.zfill(2)
            return f"{year}-{month}"
        elif re.match(r'(\w+)\s+(\d{4})', date_str, re.IGNORECASE):
            month_name, year = re.match(r'(\w+)\s+(\d{4})', date_str, re.IGNORECASE).groups()
            month = month_map.get(month_name.lower(), '01')
            return f"{year}-{month}"
        elif re.match(r'(\d{4})', date_str):
            year = date_str
            return f"{year}-01"
        return date_str

    text = re.sub(r'\b(\d{1,2}[/-]\d{4}|\d{4}[/-]\d{1,2}|\w+\s+\d{4}|\d{4})\b', standardize_date, text)
    return text.strip()

def format_date_for_display(date_str):
    """Convert YYYY-MM to YYYY-MonthName (short format, e.g., 2020-01 to 2020-Jan) or return 'N/A' if not valid."""
    if not date_str or not isinstance(date_str, str):
        return 'N/A'
    month_map = {
        '01': 'Jan', '02': 'Feb', '03': 'Mar', '04': 'Apr', '05': 'May',
        '06': 'Jun', '07': 'Jul', '08': 'Aug', '09': 'Sep', '10': 'Oct',
        '11': 'Nov', '12': 'Dec'
    }
    if re.match(r'(\d{4})-(\d{2})', date_str):
        year, month = re.match(r'(\d{4})-(\d{2})', date_str).groups()
        return f"{year}-{month_map.get(month, 'Unknown')}"
    return date_str

def has_html_formatting(text):
    if not isinstance(text, str):
        return False
    return bool(re.search(r'</?(b|i|ul|ol|li)>', text, re.IGNORECASE))

def generate_bullet_points(text, field_name):
    if not text or has_html_formatting(text):
        logging.info(f"Skipping bullet point generation for {field_name} due to existing HTML or empty text")
        return text
    prompt = f"""
    Convert the following text from the '{field_name}' field into concise bullet points. Each bullet should be a complete sentence or idea ending in a period, keeping the content professional and concise. Return only the bullet-pointed text, one bullet per line, starting with '- '.
    Text:
    {text}
    """
    try:
        response = model.generate_content(prompt)
        bullet_points = response.text.strip()
        logging.info(f"Generated bullet points for {field_name}: {bullet_points}")
        return bullet_points
    except Exception as e:
        logging.error(f"Failed to generate bullet points for {field_name}: {e}")
        return text

def extract_json(text):
    try:
        if "```json" in text:
            match = re.findall(r"```json(.*?)```", text, re.DOTALL)
            if match:
                return json.loads(match[0].strip())
        return json.loads(text)
    except Exception as e:
        logging.error(f"Error parsing JSON: {e}")
        return {}

def generate_structured_data(text):
    prompt = f"""
    You are an expert HR resume parser with advanced natural language understanding. Convert the following resume text into structured JSON with these fields:
    {{
      "name": "",
      "education_training_certifications": [{{"title": "", "start_date": "", "end_date": ""}}],
      "total_experience": "",
      "professional_summary": "",
      "netweb_projects": [{{"title": "", "description": ""}}],
      "past_projects": [{{"title": "", "description": ""}}],
      "roles_responsibilities": "",
      "technical_skills": {{
        "web_technologies": [],
        "scripting_languages": [],
        "frameworks": [],
        "databases": [],
        "web_servers": [],
        "tools": []
      }},
      "personal_details": {{
        "employee_id": "",
        "permanent_address": "",
        "local_address": "",
        "contact_number": "",
        "date_of_joining": "",
        "designation": "",
        "overall_experience": "",
        "date_of_birth": "",
        "passport_details": ""
      }},
      "work_experience": [{{"company_name": "", "start_date": "", "end_date": "", "role": "", "responsibilities": ""}}]
    }}
    Instructions:
    - Extract 'name' from the first line, prominent header, or personal details section (e.g., 'Vraj Shah').
    - Identify 'education_training_certifications' from headers like 'Education', 'Certifications', or similar. Capture degrees and certificates with start and end dates in formats like 'MM-YYYY', 'YYYY-MM', 'Month YYYY' (e.g., 'January 2020' or 'Jan 2020'), or 'YYYY'. Standardize all dates to 'YYYY-MM' format for internal storage.
    - Identify 'total_experience' from headers like 'Total Experience' or phrases indicating years of experience (e.g., '18 years').
    - Extract 'professional_summary' from headers like 'Professional Summary', 'Professional Experience Summary', or similar. Focus exclusively on tangible professional achievements (e.g., projects completed, systems implemented, awards received) and specific roles held. Exclude any mention of personal traits (e.g., 'dedicated', 'motivated', 'hardworking') or willingness-related terms (e.g., 'eager', 'willing', 'passionate'). Rephrase to emphasize measurable accomplishments and professional contributions if needed. If no clear summary exists, infer key achievements from work experience or projects, summarizing them concisely.
    - Identify 'netweb_projects' for projects explicitly mentioning 'NetWeb' or associated with the current company.
    - Identify 'past_projects' for projects under previous employers or not associated with 'NetWeb'.
    - Extract 'roles_responsibilities' from headers like 'Roles and Responsibilities', 'Key Responsibilities', or similar. If no explicit section exists, intelligently infer responsibilities from job descriptions, bullet points, or achievements under 'Work Experience', 'Professional Experience', or similar sections. Ensure responsibilities are specific, actionable tasks or outcomes (e.g., 'Developed a web application', 'Led a team of 5 engineers') and formatted as concise bullet points.
    - Extract 'technical_skills' from lists under headers like 'Technical Skill Set', 'Skills', or similar, categorizing into web_technologies, scripting_languages, frameworks, databases, web_servers, and tools.
    - For 'personal_details', extract fields like 'employee_id', 'permanent_address', etc., from sections like 'Personal Details' or similar. Standardize 'date_of_joining' and 'date_of_birth' to 'YYYY-MM'.
    - Extract 'work_experience' from sections like 'Work Experience' or 'Professional Experience', including company name, role, dates (standardized to 'YYYY-MM'), and responsibilities. Extract 'company_name' from the organization or employer name associated with each role (e.g., 'Google', 'NetWeb'). If responsibilities are missing, infer them from job descriptions or achievements in the same section.
    - Leave fields empty if data is missing, but maintain the JSON structure. Ensure all text fields are clean and concise.
    Resume:
    {text}
    """
    try:
        response = model.generate_content(prompt)
        logging.info(f"Raw AI response: {response.text}")
        data = extract_json(response.text)
        # Ensure all expected fields are present with defaults
        data = {
            'name': data.get('name', '') or '',
            'education_training_certifications': data.get('education_training_certifications', []) or [],
            'total_experience': data.get('total_experience', '') or '',
            'professional_summary': data.get('professional_summary', '') or '',
            'netweb_projects': data.get('netweb_projects', []) or [],
            'past_projects': data.get('past_projects', []) or [],
            'roles_responsibilities': data.get('roles_responsibilities', '') or '',
            'technical_skills': data.get('technical_skills', {
                'web_technologies': [], 'scripting_languages': [], 'frameworks': [],
                'databases': [], 'web_servers': [], 'tools': []
            }) or {
                'web_technologies': [], 'scripting_languages': [], 'frameworks': [],
                'databases': [], 'web_servers': [], 'tools': []
            },
            'personal_details': data.get('personal_details', {
                'employee_id': '', 'permanent_address': '', 'local_address': '',
                'contact_number': '', 'date_of_joining': '', 'designation': '',
                'overall_experience': '', 'date_of_birth': '', 'passport_details': ''
            }) or {
                'employee_id': '', 'permanent_address': '', 'local_address': '',
                'contact_number': '', 'date_of_joining': '', 'designation': '',
                'overall_experience': '', 'date_of_birth': '', 'passport_details': ''
            },
            'work_experience': data.get('work_experience', []) or []
        }
        if data['professional_summary'] and not has_html_formatting(data['professional_summary']):
            data['professional_summary'] = generate_bullet_points(data['professional_summary'], 'professional_summary')
        if data['roles_responsibilities'] and not has_html_formatting(data['roles_responsibilities']):
            data['roles_responsibilities'] = generate_bullet_points(data['roles_responsibilities'], 'roles_responsibilities')
        for exp in data['work_experience']:
            if exp.get('responsibilities') and not has_html_formatting(exp['responsibilities']):
                exp['responsibilities'] = generate_bullet_points(exp['responsibilities'], f"work_experience_responsibilities_{exp.get('role', '')}")
            # Ensure work_experience fields are not None
            exp['company_name'] = exp.get('company_name', '') or ''
            exp['start_date'] = exp.get('start_date', '') or ''
            exp['end_date'] = exp.get('end_date', '') or ''
            exp['role'] = exp.get('role', '') or ''
            exp['responsibilities'] = exp.get('responsibilities', '') or ''
        logging.info(f"Generated structured resume data: {json.dumps(data, indent=2)}")
        return data
    except Exception as e:
        logging.error(f"Failed to generate structured resume data: {e}")
        return {}

def check_grammar(text_fields):
    prompt = """
    Analyze the following text fields from a resume form and provide grammar suggestions.
    For each field, return a JSON array of suggestions with:
    [
        {
            "field": "field_name",
            "field_id": "field_id",
            "original": "original_text",
            "suggested": "corrected_text",
            "reason": "reason_for_suggestion"
        }
    ]
    Only include fields with grammar issues. Return an empty array if no suggestions are needed.
    Text fields:
    {text_fields}
    """
    try:
        text_fields_json = json.dumps(text_fields, indent=2)
        response = model.generate_content(prompt.format(text_fields=text_fields_json))
        suggestions = extract_json(response.text)
        if not isinstance(suggestions, list):
            logging.warning(f"Grammar suggestions is not a list, defaulting to []: {suggestions}")
            return []
        logging.info(f"Grammar suggestions: {json.dumps(suggestions, indent=2)}")
        return suggestions
    except Exception as e:
        logging.error(f"Failed to check grammar: {e}")
        return []

def sanitize_text(text, allow_html=False):
    if not isinstance(text, str):
        text = '' if text is None else str(text)
    if allow_html:
        return bleach.clean(text, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRIBUTES)
    return (text.replace('&', '&amp;')
               .replace('<', '&lt;')
               .replace('>', '&gt;')
               .replace('"', '&quot;')
               .replace("'", '&#x27;'))

def sanitize_profile_data(profile):
    if not isinstance(profile, dict):
        logging.warning(f"Profile is not a dict, returning default structure: {profile}")
        return {
            'name': '',
            'education_training_certifications': [],
            'total_experience': '',
            'professional_summary': '',
            'netweb_projects': [],
            'past_projects': [],
            'roles_responsibilities': '',
            'technical_skills': {
                'web_technologies': [], 'scripting_languages': [], 'frameworks': [],
                'databases': [], 'web_servers': [], 'tools': []
            },
            'personal_details': {
                'employee_id': '', 'permanent_address': '', 'local_address': '',
                'contact_number': '', 'date_of_joining': '', 'designation': '',
                'overall_experience': '', 'date_of_birth': '', 'passport_details': ''
            },
            'work_experience': []
        }

    html_fields = ['professional_summary', 'roles_responsibilities']
    sanitized = {}
    for k, v in profile.items():
        logging.debug(f"Sanitizing field {k}: {v}")
        if v is None:
            logging.warning(f"Field {k} is None, setting to default value")
            if k in ['name', 'total_experience', 'professional_summary', 'roles_responsibilities']:
                v = ''
            elif k in ['education_training_certifications', 'netweb_projects', 'past_projects', 'work_experience']:
                v = []
            elif k == 'technical_skills':
                v = {
                    'web_technologies': [], 'scripting_languages': [], 'frameworks': [],
                    'databases': [], 'web_servers': [], 'tools': []
                }
            elif k == 'personal_details':
                v = {
                    'employee_id': '', 'permanent_address': '', 'local_address': '',
                    'contact_number': '', 'date_of_joining': '', 'designation': '',
                    'overall_experience': '', 'date_of_birth': '', 'passport_details': ''
                }

        if k in html_fields:
            sanitized[k] = sanitize_text(v, allow_html=True)
        elif k == 'netweb_projects' or k == 'past_projects':
            sanitized[k] = [
                {
                    'title': sanitize_text(item.get('title', ''), allow_html=False),
                    'description': sanitize_text(item.get('description', ''), allow_html=True)
                }
                for item in v or []
            ]
        elif k == 'work_experience':
            sanitized[k] = [
                {
                    'company_name': sanitize_text(item.get('company_name', ''), allow_html=False),
                    'start_date': sanitize_text(item.get('start_date', ''), allow_html=False),
                    'end_date': sanitize_text(item.get('end_date', ''), allow_html=False),
                    'role': sanitize_text(item.get('role', ''), allow_html=False),
                    'responsibilities': sanitize_text(item.get('responsibilities', ''), allow_html=True)
                }
                for item in v or []
            ]
        elif k == 'technical_skills':
            sanitized[k] = {
                skill_type: [sanitize_text(skill, allow_html=False) for skill in (v.get(skill_type, []) or [])]
                for skill_type in ['web_technologies', 'scripting_languages', 'frameworks', 'databases', 'web_servers', 'tools']
            }
        elif k == 'personal_details':
            sanitized[k] = {
                detail_key: sanitize_text(v.get(detail_key, ''), allow_html=False)
                for detail_key in ['employee_id', 'permanent_address', 'local_address', 'contact_number',
                                  'date_of_joining', 'designation', 'overall_experience', 'date_of_birth', 'passport_details']
            }
        elif k == 'education_training_certifications':
            sanitized[k] = [
                {
                    'title': sanitize_text(item.get('title', ''), allow_html=False),
                    'start_date': sanitize_text(item.get('start_date', ''), allow_html=False),
                    'end_date': sanitize_text(item.get('end_date', ''), allow_html=False)
                }
                for item in v or []
            ]
        else:
            sanitized[k] = sanitize_text(v, allow_html=False)
    logging.debug(f"Sanitized profile: {json.dumps(sanitized, indent=2)}")
    return sanitized

def render_html_to_pdf(html_string, output_path):
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True, timeout=120000)
            try:
                page = browser.new_page()
                page.set_viewport_size({"width": 794, "height": 1123})
                page.set_content(html_string, timeout=120000, wait_until='domcontentloaded')
                page.wait_for_load_state('networkidle', timeout=120000)
                page.emulate_media(media="print")
                logging.info(f"Generating PDF at {output_path}")
                page.pdf(
                    path=output_path,
                    format="A4",
                    print_background=True,
                    margin={
                        "top": "1cm",
                        "right": "1cm",
                        "bottom": "1cm",
                        "left": "1cm"
                    },
                    prefer_css_page_size=True,
                    scale=0.8
                )
                logging.info(f"Successfully generated PDF at {output_path}")
            finally:
                browser.close()
    except Exception as e:
        logging.error(f"Failed to generate PDF: {str(e)}\n{traceback.format_exc()}")
        raise Exception(f"PDF generation failed: {str(e)}. Ensure Playwright is installed and Chromium is available.")

def cleanup_file(filepath):
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            logging.info(f"Cleaned up file: {filepath}")
    except Exception as e:
        logging.error(f"Error cleaning up file {filepath}: {str(e)}")

def should_skip_section(section_id, hidden_sections):
    return section_id in (hidden_sections or [])

def render_html_to_docx(profile, output_path, hidden_sections=None):
    if hidden_sections is None:
        hidden_sections = []
    try:
        doc = Document()
        doc.add_heading('Professional Resume', 0)
        
        if profile.get('name'):
            doc.add_heading(profile['name'], level=1)
        
        if not should_skip_section('education-section', hidden_sections) and profile.get('education_training_certifications'):
            doc.add_heading('Education, Training, and Certifications', level=1)
            for item in profile['education_training_certifications']:
                title = item.get('title', '')
                start_date = format_date_for_display(item.get('start_date', 'N/A'))
                end_date = format_date_for_display(item.get('end_date', 'N/A'))
                doc.add_paragraph(f"{title} ({start_date} - {end_date})", style='ListBullet')
        
        if not should_skip_section('experience-section', hidden_sections) and profile.get('total_experience'):
            doc.add_heading('Total Experience', level=1)
            doc.add_paragraph(profile['total_experience'])
        
        if not should_skip_section('summary-section', hidden_sections) and profile.get('professional_summary'):
            doc.add_heading('Professional Achievements', level=1)
            para = doc.add_paragraph()
            html_to_docx(para, profile['professional_summary'])
        
        if not should_skip_section('projects-section', hidden_sections):
            if profile.get('netweb_projects'):
                doc.add_heading('NetWeb Projects', level=1)
                for project in profile['netweb_projects']:
                    if project.get('title'):
                        doc.add_heading(project['title'], level=2)
                    if project.get('description'):
                        para = doc.add_paragraph()
                        html_to_docx(para, project['description'])
            
            if profile.get('past_projects'):
                doc.add_heading('Past Projects', level=1)
                for project in profile['past_projects']:
                    if project.get('title'):
                        doc.add_heading(project['title'], level=2)
                    if project.get('description'):
                        para = doc.add_paragraph()
                        html_to_docx(para, project['description'])
        
        if not should_skip_section('roles-section', hidden_sections) and profile.get('roles_responsibilities'):
            doc.add_heading('Roles and Responsibilities', level=1)
            para = doc.add_paragraph()
            html_to_docx(para, profile['roles_responsibilities'])
        
        if not should_skip_section('work-experience-section', hidden_sections) and profile.get('work_experience'):
            doc.add_heading('Work Experience', level=1)
            for exp in profile['work_experience']:
                company_name = exp.get('company_name', 'Unknown Company')
                role = exp.get('role', 'Unknown Role')
                start_date = format_date_for_display(exp.get('start_date', 'N/A'))
                end_date = format_date_for_display(exp.get('end_date', 'N/A'))
                doc.add_heading(f"{company_name} - {role} ({start_date} - {end_date})", level=2)
                if exp.get('responsibilities'):
                    para = doc.add_paragraph()
                    html_to_docx(para, exp['responsibilities'])
        
        if not should_skip_section('skills-section', hidden_sections) and profile.get('technical_skills'):
            skills = profile['technical_skills']
            if any(skills.values()):
                doc.add_heading('Technical Skills', level=1)
                for skill_type, skill_list in skills.items():
                    if skill_list:
                        doc.add_heading(skill_type.replace('_', ' ').title(), level=2)
                        for skill in skill_list:
                            doc.add_paragraph(str(skill), style='ListBullet')
        
        if not should_skip_section('personal-details-section', hidden_sections) and profile.get('personal_details'):
            personal = profile['personal_details']
            if any(personal.values()):
                doc.add_heading('Personal Details', level=1)
                for key, value in personal.items():
                    if value:
                        if key in ['date_of_joining', 'date_of_birth']:
                            value = format_date_for_display(value)
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        
        doc.save(output_path)
        logging.info(f"Generated DOCX at {output_path}")
    except Exception as e:
        logging.error(f"Error generating DOCX: {e}")
        raise

def html_to_docx(paragraph, html_text):
    from lxml import html
    try:
        tree = html.fromstring(f"<div>{html_text or ''}</div>")
        current_run = paragraph.add_run()
        
        def process_element(element, run):
            if element.tag == 'b':
                run.bold = True
                for child in element.iterchildren():
                    process_element(child, run)
            elif element.tag == 'i':
                run.italic = True
                for child in element.iterchildren():
                    process_element(child, run)
            elif element.tag == 'li':
                run.text = f"- {element.text or ''}"
                run.paragraph.style = 'ListBullet'
                for child in element.iterchildren():
                    new_run = paragraph.add_run()
                    process_element(child, new_run)
            elif element.tag in ('ul', 'ol'):
                for child in element.iterchildren():
                    process_element(child, run)
            else:
                if element.text:
                    run.text = element.text
                for child in element.iterchildren():
                    new_run = paragraph.add_run()
                    process_element(child, new_run)
        
        for child in tree.iterchildren():
            process_element(child, current_run)
    except Exception as e:
        logging.error(f"Error converting HTML to DOCX: {e}")
        paragraph.add_run(html_text or '')

def render_html_to_xlsx(profile, output_path, hidden_sections=None):
    if hidden_sections is None:
        hidden_sections = []
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Resume"
        row = 1
        
        def add_html_text(cell, text):
            if not text:
                return
            plain_text = bleach.clean(text, tags=[], strip=True).replace('\n', '; ')
            cell.value = plain_text
            apply_cell_style(cell)
        
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 50
        
        if profile.get('name'):
            ws.cell(row=row, column=1).value = "Name"
            apply_cell_style(ws.cell(row=row, column=1), is_header=True)
            ws.cell(row=row, column=2).value = profile['name']
            apply_cell_style(ws.cell(row=row, column=2))
            row += 2
        
        if not should_skip_section('education-section', hidden_sections) and profile.get('education_training_certifications'):
            ws.cell(row=row, column=1).value = "Education, Training, and Certifications"
            apply_cell_style(ws.cell(row=row, column=1), is_header=True)
            row += 1
            for item in profile['education_training_certifications']:
                start_date = format_date_for_display(item.get('start_date', 'N/A'))
                end_date = format_date_for_display(item.get('end_date', 'N/A'))
                ws.cell(row=row, column=2).value = f"{item.get('title', '')} ({start_date} - {end_date})"
                apply_cell_style(ws.cell(row=row, column=2))
                row += 1
            row += 1
        
        if not should_skip_section('experience-section', hidden_sections) and profile.get('total_experience'):
            ws.cell(row=row, column=1).value = "Total Experience"
            apply_cell_style(ws.cell(row=row, column=1), is_header=True)
            ws.cell(row=row, column=2).value = profile['total_experience']
            apply_cell_style(ws.cell(row=row, column=2))
            row += 2
        
        if not should_skip_section('summary-section', hidden_sections) and profile.get('professional_summary'):
            ws.cell(row=row, column=1).value = "Professional Achievements"
            apply_cell_style(ws.cell(row=row, column=1), is_header=True)
            add_html_text(ws.cell(row=row, column=2), profile['professional_summary'])
            row += 2
        
        if not should_skip_section('projects-section', hidden_sections):
            if profile.get('netweb_projects'):
                ws.cell(row=row, column=1).value = "NetWeb Projects"
                apply_cell_style(ws.cell(row=row, column=1), is_header=True)
                row += 1
                for project in profile['netweb_projects']:
                    if project.get('title'):
                        ws.cell(row=row, column=2).value = f"Title: {project['title']}"
                        apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                        row += 1
                    if project.get('description'):
                        ws.cell(row=row, column=2).value = "Description"
                        apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                        add_html_text(ws.cell(row=row, column=3), project['description'])
                        row += 1
                    row += 1
            
            if profile.get('past_projects'):
                ws.cell(row=row, column=1).value = "Past Projects"
                apply_cell_style(ws.cell(row=row, column=1), is_header=True)
                row += 1
                for project in profile['past_projects']:
                    if project.get('title'):
                        ws.cell(row=row, column=2).value = f"Title: {project['title']}"
                        apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                        row += 1
                    if project.get('description'):
                        ws.cell(row=row, column=2).value = "Description"
                        apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                        add_html_text(ws.cell(row=row, column=3), project['description'])
                        row += 1
                    row += 1
        
        if not should_skip_section('roles-section', hidden_sections) and profile.get('roles_responsibilities'):
            ws.cell(row=row, column=1).value = "Roles and Responsibilities"
            apply_cell_style(ws.cell(row=row, column=1), is_header=True)
            add_html_text(ws.cell(row=row, column=2), profile['roles_responsibilities'])
            row += 2
        
        if not should_skip_section('work-experience-section', hidden_sections) and profile.get('work_experience'):
            ws.cell(row=row, column=1).value = "Work Experience"
            apply_cell_style(ws.cell(row=row, column=1), is_header=True)
            row += 1
            for exp in profile['work_experience']:
                company_name = exp.get('company_name', 'Unknown Company')
                start_date = format_date_for_display(exp.get('start_date', 'N/A'))
                end_date = format_date_for_display(exp.get('end_date', 'N/A'))
                ws.cell(row=row, column=2).value = f"{company_name} - {exp.get('role', 'Unknown Role')} ({start_date} - {end_date})"
                apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                row += 1
                if exp.get('responsibilities'):
                    ws.cell(row=row, column=2).value = "Responsibilities"
                    apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                    add_html_text(ws.cell(row=row, column=3), exp['responsibilities'])
                    row += 1
                row += 1
        
        if not should_skip_section('skills-section', hidden_sections) and profile.get('technical_skills'):
            skills = profile['technical_skills']
            if any(skills.values()):
                ws.cell(row=row, column=1).value = "Technical Skills"
                apply_cell_style(ws.cell(row=row, column=1), is_header=True)
                row += 1
                for skill_type, skill_list in skills.items():
                    if skill_list:
                        ws.cell(row=row, column=2).value = skill_type.replace('_', ' ').title()
                        apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                        row += 1
                        for skill in skill_list:
                            ws.cell(row=row, column=3).value = str(skill)
                            apply_cell_style(ws.cell(row=row, column=3))
                            row += 1
                        row += 1
        
        if not should_skip_section('personal-details-section', hidden_sections) and profile.get('personal_details'):
            personal = profile['personal_details']
            if any(personal.values()):
                ws.cell(row=row, column=1).value = "Personal Details"
                apply_cell_style(ws.cell(row=row, column=1), is_header=True)
                row += 1
                for key, value in personal.items():
                    if value:
                        if key in ['date_of_joining', 'date_of_birth']:
                            value = format_date_for_display(value)
                        ws.cell(row=row, column=2).value = key.replace('_', ' ').title()
                        apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                        ws.cell(row=row, column=3).value = str(value)
                        apply_cell_style(ws.cell(row=row, column=3))
                        row += 1
                row += 1
        
        wb.save(output_path)
        logging.info(f"Generated XLSX at {output_path}")
    except Exception as e:
        logging.error(f"Error generating XLSX: {e}")
        raise

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('file_input')
        prompt_text = request.form.get('text_input')
        create_from_scratch = request.form.get('create_from_scratch')
        
        if create_from_scratch:
            logging.info("Redirecting to create_from_scratch page")
            return redirect('/create_from_scratch')
        
        text = ""
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            ext = filename.rsplit('.', 1)[1].lower()
            text = extract_text(path, ext)
            cleanup_file(path)
        elif prompt_text and prompt_text.strip():
            text = prompt_text.strip()
        if not text:
            flash("Please upload a valid file or provide resume text.")
            logging.warning("No valid input provided")
            return redirect('/')
        formatted_text = clean_formatting(text)
        profile = generate_structured_data(formatted_text)
        if not profile:
            flash("Unable to generate profile data from the provided input. Please try again.")
            logging.error("Profile generation failed")
            return redirect('/')
        session.permanent = True
        session['profile'] = sanitize_profile_data(profile)
        session['hidden_sections'] = []
        session['creation_method'] = 'upload'
        session['design'] = 'display_profile'
        logging.info(f"Initialized session: profile keys={list(profile.keys())}, hidden_sections=[], creation_method=upload, design=display_profile")
        try:
            return render_template("display_profile.html", profile=profile, hidden_sections=[])
        except Exception as e:
            logging.error(f"Template rendering error for display_profile.html: {str(e)}\n{traceback.format_exc()}")
            flash(f"Template error: {str(e)}. Please ensure display_profile.html exists.")
            return redirect('/')
    if not session.get('profile'):
        session.pop('profile', None)
        session.pop('hidden_sections', None)
        session.pop('creation_method', None)
        session.pop('design', None)
    return render_template("index.html")

@app.route('/create-from-scratch', methods=['GET'])
def create_from_scratch_hyphen():
    logging.info("Redirecting from /create-from-scratch to /create_from_scratch")
    return redirect('/create_from_scratch')

@app.route('/create_from_scratch', methods=['GET'])
def create_from_scratch():
    try:
        logging.info("Rendering create-from-scratch.html")
        return render_template("create-from-scratch.html")
    except Exception as e:
        flash(f"Template error: {str(e)}. Please ensure create-from-scratch.html exists.")
        logging.error(f"Template rendering error for create-from-scratch.html: {str(e)}")
        return redirect('/')

@app.route('/submit_from_scratch', methods=['POST'])
def submit_from_scratch():
    logging.info("Received form submission from create-from-scratch.html")
    form_data = request.form.to_dict(flat=False)
    for key, value in form_data.items():
        logging.debug(f"{key}: {value}")

    profile_data = {
        "name": request.form.get('full_name', '').strip(),
        "education_training_certifications": [
            {
                "title": item.strip() or '',
                "start_date": start.strip() or '',
                "end_date": end.strip() or ''
            }
            for item, start, end in zip(
                request.form.getlist('education_title[]'),
                request.form.getlist('education_start_date[]'),
                request.form.getlist('education_end_date[]')
            )
            if item.strip()
        ],
        "total_experience": request.form.get('total_experience', '').strip(),
        "professional_summary": request.form.get('professional_summary', '').strip(),
        "netweb_projects": [
            {"title": title.strip() or '', "description": desc.strip() or ''}
            for title, desc in zip(
                request.form.getlist('netweb_project_title[]'),
                request.form.getlist('netweb_project_description[]')
            )
            if title.strip() or desc.strip()
        ],
        "past_projects": [
            {"title": title.strip() or '', "description": desc.strip() or ''}
            for title, desc in zip(
                request.form.getlist('past_project_title[]'),
                request.form.getlist('past_project_description[]')
            )
            if title.strip() or desc.strip()
        ],
        "roles_responsibilities": request.form.get('roles_responsibilities', '').strip(),
        "technical_skills": {
            "web_technologies": [item.strip() for item in request.form.getlist('web_technologies[]') if item.strip()] or [],
            "scripting_languages": [item.strip() for item in request.form.getlist('scripting_languages[]') if item.strip()] or [],
            "frameworks": [item.strip() for item in request.form.getlist('frameworks[]') if item.strip()] or [],
            "databases": [item.strip() for item in request.form.getlist('databases[]') if item.strip()] or [],
            "web_servers": [item.strip() for item in request.form.getlist('web_servers[]') if item.strip()] or [],
            "tools": [item.strip() for item in request.form.getlist('tools[]') if item.strip()] or []
        },
        "personal_details": {
            "employee_id": request.form.get('personal_details[employee_id]', '').strip() or '',
            "permanent_address": request.form.get('personal_details[permanent_address]', '').strip() or '',
            "local_address": request.form.get('personal_details[local_address]', '').strip() or '',
            "contact_number": request.form.get('personal_details[contact_number]', '').strip() or '',
            "date_of_joining": request.form.get('personal_details[date_of_joining]', '').strip() or '',
            "designation": request.form.get('personal_details[designation]', '').strip() or '',
            "overall_experience": request.form.get('personal_details[overall_experience]', '').strip() or '',
            "date_of_birth": request.form.get('personal_details[date_of_birth]', '').strip() or '',
            "passport_details": request.form.get('personal_details[passport_details]', '').strip() or ''
        },
        "work_experience": [
            {
                "company_name": company.strip() or '',
                "start_date": start.strip() or '',
                "end_date": end.strip() or '',
                "role": role.strip() or '',
                "responsibilities": resp.strip() or ''
            }
            for company, start, end, role, resp in zip(
                request.form.getlist('work_experience[company_name][]'),
                request.form.getlist('work_experience[start_date_converted][]'),
                request.form.getlist('work_experience[end_date_converted][]'),
                request.form.getlist('work_experience[role][]'),
                request.form.getlist('work_experience[responsibilities][]')
            )
            if role.strip() or company.strip()
        ]
    }

    has_data = (
        profile_data['name'] or
        profile_data['education_training_certifications'] or
        profile_data['total_experience'] or
        profile_data['professional_summary'] or
        profile_data['netweb_projects'] or
        profile_data['past_projects'] or
        profile_data['roles_responsibilities'] or
        any(profile_data['technical_skills'].values()) or
        any(profile_data['personal_details'].values()) or
        profile_data['work_experience']
    )
    if not has_data:
        flash("No valid data was provided. Please complete at least one field and try again.")
        logging.warning("Validation failed: No valid data provided")
        return redirect('/create_from_scratch')

    session.permanent = True
    session['profile'] = sanitize_profile_data(profile_data)
    session['hidden_sections'] = []
    session['creation_method'] = 'scratch'
    session['design'] = 'display_profile'
    logging.info(f"Stored profile in session: {json.dumps(profile_data, indent=2)}")

    try:
        logging.info("Redirecting to display_profile")
        return redirect(url_for('display_profile'))
    except Exception as e:
        flash(f"Routing error: {str(e)}. Please ensure display_profile route is defined.")
        logging.error(f"Routing error for display_profile: {str(e)}")
        return redirect('/create_from_scratch')

@app.route('/check_grammar', methods=['POST'])
def check_grammar_route():
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    text_fields = {}
    fields_to_check = [
        ("name", "name", data.get('name', '') or ''),
        ("total_experience", "total_experience", data.get('total_experience', '') or ''),
        ("professional_summary", "professional_summary", bleach.clean(data.get('professional_summary', '') or '', tags=[])),
        ("roles_responsibilities", "roles_responsibilities", bleach.clean(data.get('roles_responsibilities', '') or '', tags=[]))
    ]
    
    for i, item in enumerate(data.get('education_training_certifications[]', []) or []):
        if isinstance(item, dict) and item.get('title', '').strip():
            text_fields[f"education_training_certifications[{i}]"] = {"id": f"etc_{i}", "text": item['title']}
        elif isinstance(item, str) and item.strip():
            text_fields[f"education_training_certifications[{i}]"] = {"id": f"etc_{i}", "text": item}
    
    for i, (title, desc) in enumerate(zip(data.get('netweb_projects[title][]', []) or [], data.get('netweb_projects[description][]', []) or [])):
        if title.strip():
            text_fields[f"netweb_projects[title][{i}]"] = {"id": f"netweb_title_{i}", "text": title}
        if desc.strip():
            text_fields[f"netweb_projects[description][{i}]"] = {"id": f"netweb_desc_{i}", "text": bleach.clean(desc, tags=[])}
    
    for i, (title, desc) in enumerate(zip(data.get('past_projects[title][]', []) or [], data.get('past_projects[description][]', []) or [])):
        if title.strip():
            text_fields[f"past_projects[title][{i}]"] = {"id": f"past_title_{i}", "text": title}
        if desc.strip():
            text_fields[f"past_projects[description][{i}]"] = {"id": f"past_desc_{i}", "text": bleach.clean(desc, tags=[])}
    
    for i, exp in enumerate(data.get('work_experience[]', []) or []):
        if isinstance(exp, dict):
            if exp.get('company_name', '').strip():
                text_fields[f"work_experience[company_name][{i}]"] = {"id": f"work_company_{i}", "text": exp['company_name']}
            if exp.get('role', '').strip():
                text_fields[f"work_experience[role][{i}]"] = {"id": f"work_role_{i}", "text": exp['role']}
            if exp.get('responsibilities', '').strip():
                text_fields[f"work_experience[responsibilities][{i}]"] = {"id": f"work_resp_{i}", "text": bleach.clean(exp['responsibilities'], tags=[])}
    
    for skill_type in ['web_technologies', 'scripting_languages', 'frameworks', 'databases', 'web_servers', 'tools']:
        for i, skill in enumerate(data.get(f'technical_skills[{skill_type}][]', []) or []):
            if skill.strip():
                text_fields[f"technical_skills[{skill_type}][{i}]"] = {"id": f"{skill_type.split('_')[0]}_{i}", "text": skill}
    
    for key, value in (data.get('personal_details', {}) or {}).items():
        if value.strip():
            text_fields[f"personal_details[{key}]"] = {"id": key, "text": value}
    
    suggestions = check_grammar({k: v['text'] for k, v in text_fields.items()})
    
    formatted_suggestions = [
        {
            "field": k,
            "field_id": text_fields[k]['id'],
            "original": v['original'],
            "suggested": v['suggested'],
            "reason": v['reason']
        }
        for k, v in [(k, v) for k, v in suggestions if k in text_fields]
    ]
    
    return jsonify(formatted_suggestions)

@app.route('/edit_profile')
def edit_profile():
    profile = session.get('profile', {})
    hidden_sections = session.get('hidden_sections', [])
    if not profile:
        flash("No profile data available. Please upload or create a profile.")
        logging.warning("No profile data in session for /edit_profile")
        return redirect('/')
    logging.info(f"Edit profile: profile keys={list(profile.keys())}, hidden_sections={hidden_sections}")
    try:
        return render_template("edit.html", profile=profile, hidden_sections=hidden_sections)
    except Exception as e:
        flash(f"Template error: {str(e)}. Please ensure edit.html exists.")
        logging.error(f"Template rendering error for edit.html: {str(e)}")
        return redirect('/')

@app.route('/update_profile', methods=['POST'])
def update_profile():
    action = request.form.get('action')
    logging.info("Received profile update form data:")
    form_data = request.form.to_dict(flat=False)
    for key, value in form_data.items():
        logging.debug(f"{key}: {value}")
    
    hidden_sections = []
    try:
        hidden_sections = json.loads(request.form.get('hidden_sections', '[]'))
        if not isinstance(hidden_sections, list):
            logging.warning(f"hidden_sections is not a valid list, defaulting to empty list: {hidden_sections}")
            hidden_sections = []
        hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
    except json.JSONDecodeError as e:
        logging.error(f"Error decoding hidden_sections: {e}")
        hidden_sections = []

    profile_data = {
        "name": request.form.get('name', '').strip() or '',
        "education_training_certifications": [
            {
                "title": item.strip() or '',
                "start_date": start.strip() or '',
                "end_date": end.strip() or ''
            }
            for item, start, end in zip(
                request.form.getlist('education_training_certifications[]'),
                request.form.getlist('education_start_date[]'),
                request.form.getlist('education_end_date[]')
            )
            if item.strip()
        ],
        "total_experience": request.form.get('total_experience', '').strip() or '',
        "professional_summary": request.form.get('professional_summary', '').strip() or '',
        "netweb_projects": [
            {"title": title.strip() or '', "description": desc.strip() or ''}
            for title, desc in zip(
                request.form.getlist('netweb_projects[title][]'),
                request.form.getlist('netweb_projects[description][]')
            )
            if title.strip()
        ],
        "past_projects": [
            {"title": title.strip() or '', "description": desc.strip() or ''}
            for title, desc in zip(
                request.form.getlist('past_projects[title][]'),
                request.form.getlist('past_projects[description][]')
            )
            if title.strip()
        ],
        "roles_responsibilities": request.form.get('roles_responsibilities', '').strip() or '',
        "technical_skills": {
            "web_technologies": [item.strip() for item in request.form.getlist('technical_skills[web_technologies][]') if item.strip()] or [],
            "scripting_languages": [item.strip() for item in request.form.getlist('technical_skills[scripting_languages][]') if item.strip()] or [],
            "frameworks": [item.strip() for item in request.form.getlist('technical_skills[frameworks][]') if item.strip()] or [],
            "databases": [item.strip() for item in request.form.getlist('technical_skills[databases][]') if item.strip()] or [],
            "web_servers": [item.strip() for item in request.form.getlist('technical_skills[web_servers][]') if item.strip()] or [],
            "tools": [item.strip() for item in request.form.getlist('technical_skills[tools][]') if item.strip()] or []
        },
        "personal_details": {
            "employee_id": request.form.get('personal_details[employee_id]', '').strip() or '',
            "permanent_address": request.form.get('personal_details[permanent_address]', '').strip() or '',
            "local_address": request.form.get('personal_details[local_address]', '').strip() or '',
            "contact_number": request.form.get('personal_details[contact_number]', '').strip() or '',
            "date_of_joining": request.form.get('personal_details[date_of_joining]', '').strip() or '',
            "designation": request.form.get('personal_details[designation]', '').strip() or '',
            "overall_experience": request.form.get('personal_details[overall_experience]', '').strip() or '',
            "date_of_birth": request.form.get('personal_details[date_of_birth]', '').strip() or '',
            "passport_details": request.form.get('personal_details[passport_details]', '').strip() or ''
        },
        "work_experience": [
            {
                "company_name": company.strip() or '',
                "start_date": start.strip() or '',
                "end_date": end.strip() or '',
                "role": role.strip() or '',
                "responsibilities": resp.strip() or ''
            }
            for company, start, end, role, resp in zip(
                request.form.getlist('work_experience[company_name][]'),
                request.form.getlist('work_experience[start_date_converted][]'),
                request.form.getlist('work_experience[end_date_converted][]'),
                request.form.getlist('work_experience[role][]'),
                request.form.getlist('work_experience[responsibilities][]')
            )
            if role.strip() or company.strip()
        ]
    }

    errors = []
    if not profile_data['name']:
        errors.append("Full Name is required.")
    for item in profile_data['education_training_certifications']:
        if not item['title']:
            errors.append("All education titles must be provided.")
            break
    for project in profile_data['netweb_projects'] + profile_data['past_projects']:
        if not project['title']:
            errors.append("All project titles must be provided.")
            break
    for exp in profile_data['work_experience']:
        if not exp['role']:
            errors.append("All work experience roles must be provided.")
        if not exp['company_name']:
            errors.append("All work experience company names must be provided.")
            break
    for skill_type, skills in profile_data['technical_skills'].items():
        for skill in skills:
            if not skill:
                errors.append(f"All {skill_type.replace('_', ' ')} skills must be provided.")
                break

    has_data = (
        profile_data['name'] or
        profile_data['education_training_certifications'] or
        profile_data['total_experience'] or
        profile_data['professional_summary'] or
        profile_data['netweb_projects'] or
        profile_data['past_projects'] or
        profile_data['roles_responsibilities'] or
        any(profile_data['technical_skills'].values()) or
        any(profile_data['personal_details'].values()) or
        profile_data['work_experience']
    )
    if not has_data:
        errors.append("No valid data was provided. Please complete at least one field.")

    if errors:
        for error in errors:
            flash(error)
        logging.warning(f"Validation failed: {errors}")
        return render_template("edit.html", profile=profile_data, hidden_sections=hidden_sections)

    profile_data = sanitize_profile_data(profile_data)
    session.permanent = True
    session['profile'] = profile_data
    session['hidden_sections'] = hidden_sections
    logging.info(f"Updated session profile: {json.dumps(profile_data, indent=2)}")
    logging.info(f"Updated session hidden_sections: {hidden_sections}")

    if action == 'save':
        logging.info(f"Rendering {session.get('design', 'display_profile')}.html with hidden_sections: {hidden_sections}")
        try:
            return render_template(f"{session.get('design', 'display_profile')}.html", profile=profile_data, hidden_sections=hidden_sections)
        except Exception as e:
            flash(f"Template error: {str(e)}. Please ensure {session.get('design', 'display_profile')}.html exists.")
            logging.error(f"Template rendering error for {session.get('design', 'display_profile')}.html: {str(e)}\n{traceback.format_exc()}")
            return redirect('/')
    else:
        flash("Invalid action requested. Please save or update the profile.")
        logging.warning("Invalid action received")
        return render_template("edit.html", profile=profile_data, hidden_sections=hidden_sections)

@app.route('/download', methods=['POST'])
def download():
    try:
        profile = session.get('profile', {})
        logging.info(f"Download PDF - Session profile: {json.dumps(profile, indent=2) if profile else 'None'}, creation_method: {session.get('creation_method', 'None')}")

        if not profile:
            flash("No profile data available. Please create or upload a profile.")
            logging.warning("No profile data in session for /download")
            return redirect('/')

        has_data = (
            profile.get('name') or
            profile.get('education_training_certifications') or
            profile.get('total_experience') or
            profile.get('professional_summary') or
            profile.get('netweb_projects') or
            profile.get('past_projects') or
            profile.get('roles_responsibilities') or
            any(profile.get('technical_skills', {}).values()) or
            any(profile.get('personal_details', {}).values()) or
            profile.get('work_experience')
        )

        if not has_data:
            flash("Profile data is empty or invalid. Please complete at least one field.")
            logging.warning("Profile data is empty or invalid")
            return redirect('/')

        hidden_sections = []
        try:
            hidden_sections_str = request.form.get('hidden_sections', '[]')
            logging.debug(f"Raw hidden_sections from form: {hidden_sections_str}")
            hidden_sections = json.loads(hidden_sections_str)
            if not isinstance(hidden_sections, list):
                logging.warning(f"hidden_sections is not a list, defaulting to empty: {hidden_sections}")
                hidden_sections = []
            hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
        except json.JSONDecodeError as e:
            logging.error(f"Error decoding hidden_sections: {e}")
            hidden_sections = []

        logging.info(f"Download PDF - hidden_sections: {hidden_sections}")

        session.permanent = True
        session['hidden_sections'] = hidden_sections

        safe_profile = sanitize_profile_data(profile)

        design = session.get('design', 'display_profile')
        template = f"{design}.html"

        try:
            html_content = render_template(
                template,
                profile=safe_profile,
                hidden_sections=hidden_sections
            )
        except Exception as e:
            logging.error(f"Template rendering failed for {template}: {str(e)}\n{traceback.format_exc()}")
            flash(f"Failed to render resume template: {str(e)}")
            return redirect('/')

        output_path = os.path.join(GENERATED_FOLDER, f"resume_{uuid.uuid4().hex}.pdf")

        try:
            render_html_to_pdf(html_content, output_path)
        except Exception as e:
            logging.error(f"PDF generation failed: {str(e)}\n{traceback.format_exc()}")
            flash(f"Failed to generate PDF: {str(e)}")
            return redirect('/')

        if not os.path.exists(output_path):
            logging.error(f"PDF file not found at {output_path}")
            flash("PDF file could not be generated. Please try again.")
            return redirect('/')

        if os.path.getsize(output_path) == 0:
            logging.error(f"Generated PDF at {output_path} is empty")
            flash("Generated PDF is empty. Please try again.")
            cleanup_file(output_path)
            return redirect('/')

        safe_name = safe_profile.get('name', 'Resume').replace(' ', '_').replace('/', '_')
        download_name = f"{safe_name}_Resume.pdf"

        logging.info(f"Successfully generated PDF at {output_path}, sending as {download_name}")

        try:
            return send_file(
                output_path,
                as_attachment=True,
                download_name=download_name,
                mimetype='application/pdf'
            )
        except Exception as e:
            logging.error(f"Failed to send PDF file: {str(e)}\n{traceback.format_exc()}")
            flash(f"Failed to download PDF: {str(e)}")
            return redirect('/')
        finally:
            cleanup_file(output_path)

    except Exception as e:
        logging.error(f"Unexpected error in download route: {str(e)}\n{traceback.format_exc()}")
        flash(f"An unexpected error occurred: {str(e)}")
        return redirect('/')

@app.route('/download_docx', methods=['POST'])
def download_docx():
    profile = session.get('profile', {})
    logging.info(f"Download DOCX - Session profile keys: {list(profile.keys()) if profile else 'None'}")
    logging.info(f"Download DOCX - Form data: {request.form.to_dict()}")

    hidden_sections = []
    try:
        hidden_sections_str = request.form.get('hidden_sections', '[]')
        hidden_sections = json.loads(hidden_sections_str)
        if not isinstance(hidden_sections, list):
            logging.warning(f"hidden_sections is not a valid list, defaulting to empty list: {hidden_sections}")
            hidden_sections = []
        hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
    except json.JSONDecodeError as e:
        logging.error(f"Error decoding hidden_sections: {e}\n{traceback.format_exc()}")
        hidden_sections = []

    logging.info(f"Download DOCX request - hidden_sections: {hidden_sections}")

    if not profile:
        flash("No profile data available. Please create or upload a profile.")
        logging.warning("No profile data in session for /download_docx")
        return redirect('/')

    output_path = None
    try:
        session.permanent = True
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.docx")
        render_html_to_docx(profile, output_path, hidden_sections)
        if not os.path.exists(output_path):
            logging.error(f"DOCX file not found at {output_path}")
            flash("DOCX file could not be generated. Please try again.")
            return redirect('/')
        return send_file(
            output_path,
            as_attachment=True,
            download_name=f"{profile.get('name', 'Employee').replace(' ', '_')}_Profile.docx"
        )
    except Exception as e:
        flash(f"Failed to generate DOCX: {str(e)}. Please try again or contact support.")
        logging.error(f"DOCX generation failed: {str(e)}\n{traceback.format_exc()}")
        return redirect('/')
    finally:
        if output_path:
            cleanup_file(output_path)

@app.route('/download_xlsx', methods=['POST'])
def download_xlsx():
    profile = session.get('profile', {})
    logging.info(f"Download XLSX - Session profile keys: {list(profile.keys()) if profile else 'None'}")
    logging.info(f"Download XLSX - Form data: {request.form.to_dict()}")

    hidden_sections = []
    try:
        hidden_sections_str = request.form.get('hidden_sections', '[]')
        hidden_sections = json.loads(hidden_sections_str)
        if not isinstance(hidden_sections, list):
            logging.warning(f"hidden_sections is not a valid list, defaulting to empty list: {hidden_sections}")
            hidden_sections = []
        hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
    except json.JSONDecodeError as e:
        logging.error(f"Error decoding hidden_sections: {e}\n{traceback.format_exc()}")
        hidden_sections = []

    logging.info(f"Download XLSX request - hidden_sections: {hidden_sections}")

    if not profile:
        flash("No profile data available. Please create or upload a profile.")
        logging.warning("No profile data in session for /download_xlsx")
        return redirect('/')

    skills_only = request.form.get('skills_only') == 'true'
    output_path = None
    try:
        session.permanent = True
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.xlsx")
        
        if skills_only:
            wb = Workbook()
            ws = wb.active
            ws.title = "Technical Skills"
            row = 1
            
            ws.column_dimensions['A'].width = 30
            ws.column_dimensions['B'].width = 50
            
            skills = profile.get('technical_skills', {})
            if any(skills.values()) and 'skills-section' not in hidden_sections:
                ws.cell(row=row, column=1).value = "Technical Skills"
                apply_cell_style(ws.cell(row=row, column=1), is_header=True)
                row += 1
                for skill_type, skill_list in skills.items():
                    if skill_list:
                        ws.cell(row=row, column=2).value = skill_type.replace('_', ' ').title()
                        apply_cell_style(ws.cell(row=row, column=2), is_header=True)
                        row += 1
                        for skill in skill_list:
                            ws.cell(row=row, column=2).value = str(skill)
                            apply_cell_style(ws.cell(row=row, column=2))
                            row += 1
                        row += 1
            
            wb.save(output_path)
            logging.info(f"Generated skills-only XLSX at {output_path}")
        else:
            render_html_to_xlsx(profile, output_path, hidden_sections)

        if not os.path.exists(output_path):
            logging.error(f"XLSX file not found at {output_path}")
            flash("XLSX file could not be generated. Please try again.")
            return redirect('/')
        
        download_name = f"{profile.get('name', 'Employee').replace(' ', '_')}_Skills.xlsx" if skills_only else f"{profile.get('name', 'Employee').replace(' ', '_')}_Profile.xlsx"
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=download_name
        )
    except Exception as e:
        flash(f"Failed to generate XLSX: {str(e)}. Please try again or contact support.")
        logging.error(f"XLSX generation failed: {str(e)}\n{traceback.format_exc()}")
        return redirect('/')
    finally:
        if output_path:
            cleanup_file(output_path)

@app.route('/display_profile')
def display_profile():
    profile = session.get('profile', {})
    hidden_sections = session.get('hidden_sections', [])
    design = session.get('design', 'display_profile')
    
    if not profile:
        flash("No profile data available. Please upload or create a profile.")
        logging.warning("No profile data in session for /display_profile")
        return redirect('/')
    
    # Log the entire profile to check for None values
    logging.debug(f"Rendering {design}.html with profile: {json.dumps(profile, indent=2)}")
    logging.debug(f"Hidden sections: {hidden_sections}")
    
    # Sanitize profile to ensure no None values
    safe_profile = sanitize_profile_data(profile)
    
    try:
        return render_template(f"{design}.html", profile=safe_profile, hidden_sections=hidden_sections)
    except Exception as e:
        logging.error(f"Template rendering error for {design}.html: {str(e)}\n{traceback.format_exc()}")
        flash(f"Template error: {str(e)}. Please ensure {design}.html exists.")
        return redirect('/')

@app.route('/switch_design', methods=['POST'])
def switch_design():
    try:
        design = request.form.get('design')
        if design not in ['display_profile', 'd1', 'd2', 'd3']:
            logging.error(f"Invalid design selected: {design}")
            return jsonify({'error': 'Invalid design selected'}), 400

        session.permanent = True
        session['design'] = design
        logging.info(f"Switched design to: {design}")

        profile = session.get('profile', {})
        hidden_sections = session.get('hidden_sections', [])

        if not profile:
            logging.warning("No profile data in session for /switch_design")
            return jsonify({'error': 'No profile data available'}), 400

        safe_profile = sanitize_profile_data(profile)
        logging.debug(f"Switch design - profile: {json.dumps(safe_profile, indent=2)}")
        logging.debug(f"Switch design - hidden_sections: {hidden_sections}")

        try:
            html = render_template(f'{design}.html', profile=safe_profile, hidden_sections=hidden_sections)
            logging.info(f"Successfully rendered {design}.html")
            return jsonify({'html': html})
        except Exception as e:
            logging.error(f"Template rendering failed for {design}.html: {str(e)}\n{traceback.format_exc()}")
            return jsonify({'error': f"Failed to render template: {str(e)}"}), 500

    except Exception as e:
        logging.error(f"Unexpected error in switch_design: {str(e)}\n{traceback.format_exc()}")
        return jsonify({'error': f"An unexpected error occurred: {str(e)}"}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
