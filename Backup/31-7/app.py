import os
import uuid
import json
import re
from flask import Flask, render_template, request, send_file, flash, redirect, session, jsonify
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import Workbook
import google.generativeai as genai
from playwright.sync_api import sync_playwright

# Gemini API Key
genai.configure(api_key="AIzaSyDzzIncgM-mfsad8QYWyYG5PvQtXYdlpbs")

app = Flask(__name__)
app.secret_key = 'supersecret'
UPLOAD_FOLDER = 'Uploads'
GENERATED_FOLDER = 'generated'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

model = genai.GenerativeModel("gemini-2.5-flash")

# Valid section IDs for validation
VALID_SECTION_IDS = [
    'education-section',
    'experience-section',
    'summary-section',
    'projects-section',
    'roles-section',
    'skills-section',
    'personal-details-section'
]

# === Helper Functions ===
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(filepath, ext):
    try:
        if ext == 'pdf':
            with open(filepath, 'rb') as f:
                text = ' '.join([p.extract_text() for p in PdfReader(f).pages if p.extract_text()])
                print(f"Extracted text from PDF file '{filepath}': {text[:100]}...")
                return text
        elif ext == 'docx':
            text = '\n'.join([p.text for p in Document(filepath).paragraphs])
            print(f"Extracted text from DOCX file '{filepath}': {text[:100]}...")
            return text
        elif ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
                print(f"Extracted text from TXT file '{filepath}': {text[:100]}...")
                return text
        return ""
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""

def clean_formatting(text):
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'•', '-', text)
    text = re.sub(r'\n\s*-\s*', '\n- ', text)
    return text.strip()

def extract_json(text):
    try:
        if "```json" in text:
            match = re.findall(r"```json(.*?)```", text, re.DOTALL)
            if match:
                return json.loads(match[0].strip())
        return json.loads(text)
    except Exception as e:
        print(f"Error parsing JSON: {e}")
        return {}

def generate_structured_data(text):
    prompt = f"""
    You are an expert HR resume parser. Convert the following text into structured JSON with these fields:
    {{
      "name": "",
      "education_training_certifications": [],
      "total_experience": "",
      "professional_summary": "",
      "netweb_projects": [{{"title": "", "description": ""}}],
      "past_projects": [{{"title": "", "description": ""}}],
      "roles_responsibilities": "",
      "technical_skills": {{
        "web_technologies": [],
        "scripting_languages": [],
        "frameworks": [],
        "databases": [],
        "web_servers": [],
        "tools": []
      }},
      "personal_details": {{
        "employee_id": "",
        "permanent_address": "",
        "local_address": "",
        "contact_number": "",
        "date_of_joining": "",
        "designation": "",
        "overall_experience": "",
        "date_of_birth": "",
        "passport_details": ""
      }}
    }}
    Return ONLY valid JSON. Do not return markdown or explanation.
    Resume:
    {text}
    """
    try:
        response = model.generate_content(prompt)
        data = extract_json(response.text)
        print(f"Generated structured resume data: {json.dumps(data, indent=2)}")
        return data
    except Exception as e:
        print(f"Failed to generate structured resume data: {e}")
        return {}

def check_grammar(text_fields):
    prompt = """
    You are an expert grammar checker. Analyze the following text fields from a resume form and provide grammar suggestions.
    For each field, return a JSON array of suggestions with the following structure:
    [
        {
            "field": "field_name",
            "field_id": "field_id",
            "original": "original_text",
            "suggested": "corrected_text",
            "reason": "reason_for_suggestion"
        }
    ]
    Only include fields that have grammar issues. Skip fields with no issues. Return an empty array if no suggestions are needed.
    Text fields:
    {text_fields}
    """
    try:
        # Prepare text fields for the prompt
        text_fields_json = json.dumps(text_fields, indent=2)
        response = model.generate_content(prompt.format(text_fields=text_fields_json))
        suggestions = extract_json(response.text)
        if not isinstance(suggestions, list):
            print(f"Warning: Grammar suggestions is not a list, defaulting to []: {suggestions}")
            return []
        print(f"Grammar suggestions: {json.dumps(suggestions, indent=2)}")
        return suggestions
    except Exception as e:
        print(f"Failed to check grammar: {e}")
        return []

def render_html_to_pdf(html_string, output_path):
    try:
        with sync_playwright() as p:
            browser = None
            try:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.set_content(html_string)
                page.pdf(path=output_path, format="A4", print_background=True)
                print(f"Successfully generated PDF at {output_path}")
            except Exception as e:
                print(f"Playwright error during PDF generation: {str(e)}")
                raise
            finally:
                if browser:
                    browser.close()
    except Exception as e:
        print(f"Failed to generate PDF: {str(e)}")
        raise Exception(f"PDF generation failed: {str(e)}")

def cleanup_file(filepath):
    """Remove a file if it exists."""
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            print(f"Cleaned up file: {filepath}")
    except Exception as e:
        print(f"Error cleaning up file {filepath}: {str(e)}")

def should_skip_section(section_id, hidden_sections):
    """Determine if a section should be excluded based on the list of hidden section IDs."""
    return section_id in hidden_sections

def render_html_to_docx(profile, output_path, hidden_sections=None):
    if hidden_sections is None:
        hidden_sections = []
    
    try:
        doc = Document()
        doc.add_heading('Professional Resume', 0)
        
        # Name
        if profile.get('name'):
            doc.add_heading('Name', level=1)
            doc.add_paragraph(profile['name'])
        
        # Education
        if not should_skip_section('education-section', hidden_sections) and profile.get('education_training_certifications'):
            doc.add_heading('Education, Training, and Certifications', level=1)
            for item in profile['education_training_certifications']:
                doc.add_paragraph(str(item), style='ListBullet')
        
        # Total Experience
        if not should_skip_section('experience-section', hidden_sections) and profile.get('total_experience'):
            doc.add_heading('Total Experience', level=1)
            doc.add_paragraph(profile['total_experience'])
        
        # Professional Summary
        if not should_skip_section('summary-section', hidden_sections) and profile.get('professional_summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(profile['professional_summary'])
        
        # NetWeb Projects
        if not should_skip_section('projects-section', hidden_sections) and profile.get('netweb_projects'):
            doc.add_heading('NetWeb Projects', level=1)
            for project in profile['netweb_projects']:
                if project.get('title'):
                    doc.add_heading(project['title'], level=2)
                if project.get('description'):
                    doc.add_paragraph(project['description'])
        
        # Past Projects
        if not should_skip_section('projects-section', hidden_sections) and profile.get('past_projects'):
            doc.add_heading('Past Projects', level=1)
            for project in profile['past_projects']:
                if project.get('title'):
                    doc.add_heading(project['title'], level=2)
                if project.get('description'):
                    doc.add_paragraph(project['description'])
        
        # Roles and Responsibilities
        if not should_skip_section('roles-section', hidden_sections) and profile.get('roles_responsibilities'):
            doc.add_heading('Roles and Responsibilities', level=1)
            doc.add_paragraph(profile['roles_responsibilities'])
        
        # Technical Skills
        if not should_skip_section('skills-section', hidden_sections) and profile.get('technical_skills'):
            skills = profile['technical_skills']
            if any(skills.values()):
                doc.add_heading('Technical Skills', level=1)
                for skill_type, skill_list in skills.items():
                    if skill_list:
                        doc.add_heading(skill_type.replace('_', ' ').title(), level=2)
                        for skill in skill_list:
                            doc.add_paragraph(str(skill), style='ListBullet')
        
        # Personal Details
        if not should_skip_section('personal-details-section', hidden_sections) and profile.get('personal_details'):
            personal = profile['personal_details']
            if any(personal.values()):
                doc.add_heading('Personal Details', level=1)
                for key, value in personal.items():
                    if value:
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
        
        doc.save(output_path)
    except Exception as e:
        print(f"Error generating DOCX: {e}")

def render_html_to_xlsx(profile, output_path, hidden_sections=None):
    if hidden_sections is None:
        hidden_sections = []
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Resume"
        row = 1
        
        # Name
        if profile.get('name'):
            ws.cell(row=row, column=1).value = "Name"
            ws.cell(row=row, column=2).value = profile['name']
            row += 1
        
        # Education
        if not should_skip_section('education-section', hidden_sections) and profile.get('education_training_certifications'):
            ws.cell(row=row, column=1).value = "Education, Training, and Certifications"
            row += 1
            for item in profile['education_training_certifications']:
                ws.cell(row=row, column=2).value = str(item)
                row += 1
        
        # Total Experience
        if not should_skip_section('experience-section', hidden_sections) and profile.get('total_experience'):
            ws.cell(row=row, column=1).value = "Total Experience"
            ws.cell(row=row, column=2).value = profile['total_experience']
            row += 1
        
        # Professional Summary
        if not should_skip_section('summary-section', hidden_sections) and profile.get('professional_summary'):
            ws.cell(row=row, column=1).value = "Professional Summary"
            ws.cell(row=row, column=2).value = profile['professional_summary']
            row += 1
        
        # NetWeb Projects
        if not should_skip_section('projects-section', hidden_sections) and profile.get('netweb_projects'):
            ws.cell(row=row, column=1).value = "NetWeb Projects"
            row += 1
            for project in profile['netweb_projects']:
                if project.get('title'):
                    ws.cell(row=row, column=2).value = f"Title: {project['title']}"
                    row += 1
                if project.get('description'):
                    ws.cell(row=row, column=2).value = f"Description: {project['description']}"
                    row += 1
        
        # Past Projects
        if not should_skip_section('projects-section', hidden_sections) and profile.get('past_projects'):
            ws.cell(row=row, column=1).value = "Past Projects"
            row += 1
            for project in profile['past_projects']:
                if project.get('title'):
                    ws.cell(row=row, column=2).value = f"Title: {project['title']}"
                    row += 1
                if project.get('description'):
                    ws.cell(row=row, column=2).value = f"Description: {project['description']}"
                    row += 1
        
        # Roles and Responsibilities
        if not should_skip_section('roles-section', hidden_sections) and profile.get('roles_responsibilities'):
            ws.cell(row=row, column=1).value = "Roles and Responsibilities"
            ws.cell(row=row, column=2).value = profile['roles_responsibilities']
            row += 1
        
        # Technical Skills
        if not should_skip_section('skills-section', hidden_sections) and profile.get('technical_skills'):
            skills = profile['technical_skills']
            if any(skills.values()):
                ws.cell(row=row, column=1).value = "Technical Skills"
                row += 1
                for skill_type, skill_list in skills.items():
                    if skill_list:
                        ws.cell(row=row, column=2).value = skill_type.replace('_', ' ').title()
                        row += 1
                        for skill in skill_list:
                            ws.cell(row=row, column=3).value = str(skill)
                            row += 1
        
        # Personal Details
        if not should_skip_section('personal-details-section', hidden_sections) and profile.get('personal_details'):
            personal = profile['personal_details']
            if any(personal.values()):
                ws.cell(row=row, column=1).value = "Personal Details"
                row += 1
                for key, value in personal.items():
                    if value:
                        ws.cell(row=row, column=2).value = key.replace('_', ' ').title()
                        ws.cell(row=row, column=3).value = str(value)
                        row += 1
        
        wb.save(output_path)
    except Exception as e:
        print(f"Error generating XLSX: {e}")

# === Routes ===
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('file_input')
        prompt_text = request.form.get('text_input')
        text = ""
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            ext = filename.rsplit('.', 1)[1].lower()
            text = extract_text(path, ext)
        elif prompt_text and prompt_text.strip():
            text = prompt_text.strip()
        if not text:
            flash("Please upload a valid file or provide resume text.")
            return redirect('/')
        formatted_text = clean_formatting(text)
        profile = generate_structured_data(formatted_text)
        if not profile:
            flash("Unable to generate profile data from the provided input. Please try again.")
            return redirect('/')
        session['profile'] = profile
        session['hidden_sections'] = []
        print(f"Initialized session: profile={json.dumps(profile, indent=2)}, hidden_sections=[]")
        return render_template("profile.html", profile=profile, hidden_sections=[])
    session.pop('profile', None)
    session.pop('hidden_sections', None)
    return render_template("index.html")

@app.route('/check_grammar', methods=['POST'])
def check_grammar_route():
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
    
    # Collect text fields for grammar checking
    text_fields = {}
    fields_to_check = [
        ("name", "name", data.get('name', '')),
        ("total_experience", "total_experience", data.get('total_experience', '')),
        ("professional_summary", "professional_summary", data.get('professional_summary', '')),
        ("roles_responsibilities", "roles_responsibilities", data.get('roles_responsibilities', ''))
    ]
    
    # Add education_training_certifications
    for i, item in enumerate(data.get('education_training_certifications[]', [])):
        if item.strip():
            text_fields[f"education_training_certifications[{i}]"] = {"id": f"etc_{i}", "text": item}
    
    # Add netweb_projects
    for i, (title, desc) in enumerate(zip(data.get('netweb_projects[title][]', []), data.get('netweb_projects[description][]', []))):
        if title.strip():
            text_fields[f"netweb_projects[title][{i}]"] = {"id": f"netweb_title_{i}", "text": title}
        if desc.strip():
            text_fields[f"netweb_projects[description][{i}]"] = {"id": f"netweb_desc_{i}", "text": desc}
    
    # Add past_projects
    for i, (title, desc) in enumerate(zip(data.get('past_projects[title][]', []), data.get('past_projects[description][]', []))):
        if title.strip():
            text_fields[f"past_projects[title][{i}]"] = {"id": f"past_title_{i}", "text": title}
        if desc.strip():
            text_fields[f"past_projects[description][{i}]"] = {"id": f"past_desc_{i}", "text": desc}
    
    # Add technical_skills
    for skill_type in ['web_technologies', 'scripting_languages', 'frameworks', 'databases', 'web_servers', 'tools']:
        for i, skill in enumerate(data.get(f'technical_skills[{skill_type}][]', [])):
            if skill.strip():
                text_fields[f"technical_skills[{skill_type}][{i}]"] = {"id": f"{skill_type.split('_')[0]}_{i}", "text": skill}
    
    # Add personal_details
    for key, value in data.get('personal_details', {}).items():
        if value.strip():
            text_fields[f"personal_details[{key}]"] = {"id": key, "text": value}
    
    # Check grammar
    suggestions = check_grammar({k: v['text'] for k, v in text_fields.items()})
    
    # Format suggestions with field IDs
    formatted_suggestions = [
        {
            "field": k,
            "field_id": text_fields[k]['id'],
            "original": v['original'],
            "suggested": v['suggested'],
            "reason": v['reason']
        }
        for k, v in [(k, v) for k, v in suggestions if k in text_fields]
    ]
    
    return jsonify(formatted_suggestions)

@app.route('/update_profile', methods=['POST'])
def update_profile():
    action = request.form.get('action')
    print("Received profile update form data:")
    form_data = request.form.to_dict(flat=False)
    for key, value in form_data.items():
        print(f"{key}: {value}")
    
    # Parse hidden sections
    hidden_sections = []
    try:
        hidden_sections = json.loads(request.form.get('hidden_sections', '[]'))
        if not isinstance(hidden_sections, list):
            print(f"Warning: hidden_sections is not a valid list, defaulting to empty list: {hidden_sections}")
            hidden_sections = []
        # Validate section IDs
        hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
    except json.JSONDecodeError as e:
        print(f"Error decoding hidden_sections: {e}, defaulting to []")
        hidden_sections = []
    print(f"Parsed hidden_sections: {hidden_sections}")

    profile_data = {
        "name": request.form.get('name', '').strip(),
        "education_training_certifications": [item.strip() for item in request.form.getlist('education_training_certifications[]') if item.strip()],
        "total_experience": request.form.get('total_experience', '').strip(),
        "professional_summary": request.form.get('professional_summary', '').strip(),
        "netweb_projects": [
            {"title": title.strip(), "description": desc.strip()}
            for title, desc in zip(
                request.form.getlist('netweb_projects[title][]'),
                request.form.getlist('netweb_projects[description][]')
            )
            if title.strip() or desc.strip()
        ],
        "past_projects": [
            {"title": title.strip(), "description": desc.strip()}
            for title, desc in zip(
                request.form.getlist('past_projects[title][]'),
                request.form.getlist('past_projects[description][]')
            )
            if title.strip() or desc.strip()
        ],
        "roles_responsibilities": request.form.get('roles_responsibilities', '').strip(),
        "technical_skills": {
            "web_technologies": [item.strip() for item in request.form.getlist('technical_skills[web_technologies][]') if item.strip()],
            "scripting_languages": [item.strip() for item in request.form.getlist('technical_skills[scripting_languages][]') if item.strip()],
            "frameworks": [item.strip() for item in request.form.getlist('technical_skills[frameworks][]') if item.strip()],
            "databases": [item.strip() for item in request.form.getlist('technical_skills[databases][]') if item.strip()],
            "web_servers": [item.strip() for item in request.form.getlist('technical_skills[web_servers][]') if item.strip()],
            "tools": [item.strip() for item in request.form.getlist('technical_skills[tools][]') if item.strip()]
        },
        "personal_details": {
            "employee_id": request.form.get('personal_details[employee_id]', '').strip(),
            "permanent_address": request.form.get('personal_details[permanent_address]', '').strip(),
            "local_address": request.form.get('personal_details[local_address]', '').strip(),
            "contact_number": request.form.get('personal_details[contact_number]', '').strip(),
            "date_of_joining": request.form.get('personal_details[date_of_joining]', '').strip(),
            "designation": request.form.get('personal_details[designation]', '').strip(),
            "overall_experience": request.form.get('personal_details[overall_experience]', '').strip(),
            "date_of_birth": request.form.get('personal_details[date_of_birth]', '').strip(),
            "passport_details": request.form.get('personal_details[passport_details]', '').strip()
        }
    }

    # Validate profile data
    has_data = (
        profile_data['name'] or
        profile_data['education_training_certifications'] or
        profile_data['total_experience'] or
        profile_data['professional_summary'] or
        profile_data['netweb_projects'] or
        profile_data['past_projects'] or
        profile_data['roles_responsibilities'] or
        any(profile_data['technical_skills'].values()) or
        any(profile_data['personal_details'].values())
    )
    if not has_data:
        flash("No valid data was provided. Please complete at least one field and try again.")
        print("Validation failed: No valid data provided")
        return render_template("profile.html", profile=profile_data, hidden_sections=hidden_sections)

    session['profile'] = profile_data
    session['hidden_sections'] = hidden_sections
    print(f"Updated session profile: {json.dumps(profile_data, indent=2)}")
    print(f"Updated session hidden_sections: {hidden_sections}")

    if action == 'save':
        print(f"Rendering display_profile.html with hidden_sections: {hidden_sections}")
        return render_template("display_profile.html", profile=profile_data, hidden_sections=hidden_sections)
    else:
        flash("Invalid action requested. Please save or update the profile.")
        print("Invalid action received")
        return render_template("profile.html", profile=profile_data, hidden_sections=hidden_sections)

@app.route('/edit_profile')
def edit_profile():
    profile = session.get('profile', {})
    hidden_sections = session.get('hidden_sections', [])
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /edit_profile")
        return redirect('/')
    print(f"Edit profile: profile={json.dumps(profile, indent=2)}, hidden_sections={hidden_sections}")
    return render_template("profile.html", profile=profile, hidden_sections=hidden_sections)

@app.route('/download', methods=['POST'])
def download():
    # Retrieve profile from session
    profile = session.get('profile', {})
    if not profile:
        flash("No profile data available. Please create or upload a profile.")
        print("No profile data in session for /download")
        return redirect('/')
    
    # Validate profile data
    has_data = (
        profile.get('name') or
        profile.get('education_training_certifications') or
        profile.get('total_experience') or
        profile.get('professional_summary') or
        profile.get('netweb_projects') or
        profile.get('past_projects') or
        profile.get('roles_responsibilities') or
        any(profile.get('technical_skills', {}).values()) or
        any(profile.get('personal_details', {}).values())
    )
    if not has_data:
        flash("Profile data is empty or invalid. Please complete at least one field.")
        print("Profile data is empty or invalid")
        return redirect('/')

    # Parse hidden sections from form data
    hidden_sections = []
    try:
        hidden_sections_str = request.form.get('hidden_sections', '[]')
        hidden_sections = json.loads(hidden_sections_str)
        if not isinstance(hidden_sections, list):
            print(f"Warning: hidden_sections is not a valid list, defaulting to empty list: {hidden_sections}")
            hidden_sections = []
        # Validate section IDs
        hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
    except json.JSONDecodeError as e:
        print(f"Error decoding hidden_sections: {e}, defaulting to []")
        hidden_sections = []
    
    print(f"Download request - profile keys: {list(profile.keys())}, hidden_sections: {hidden_sections}")

    try:
        # Update session with current hidden sections
        session['hidden_sections'] = hidden_sections
        
        # Render HTML with hidden sections
        html = render_template("display_profile.html", profile=profile, hidden_sections=hidden_sections)
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.pdf")
        
        # Generate PDF
        render_html_to_pdf(html, output_path)
        
        # Send file and clean up
        try:
            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"{profile.get('name', 'Employee')}_Profile.pdf"
            )
        finally:
            cleanup_file(output_path)
    except Exception as e:
        flash(f"Failed to generate PDF: {str(e)}. Please try again or contact support.")
        print(f"PDF generation failed: {str(e)}")
        return redirect('/')

@app.route('/download_docx', methods=['POST'])
def download_docx():
    profile = session.get('profile', {})
    
    # Get hidden sections from form data
    hidden_sections = []
    try:
        hidden_sections_str = request.form.get('hidden_sections', '[]')
        hidden_sections = json.loads(hidden_sections_str)
        if not isinstance(hidden_sections, list):
            print(f"Warning: hidden_sections is not a valid list, defaulting to empty list: {hidden_sections}")
            hidden_sections = []
        # Validate section IDs
        hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
    except json.JSONDecodeError as e:
        print(f"Error decoding hidden_sections: {e}, defaulting to []")
        hidden_sections = []
    
    print(f"Download DOCX request - hidden_sections: {hidden_sections}")
    
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /download_docx")
        return redirect('/')
    try:
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.docx")
        render_html_to_docx(profile, output_path, hidden_sections)
        print(f"Generated DOCX: {output_path} with hidden sections: {hidden_sections}")
        return send_file(output_path, as_attachment=True, download_name="Employee_Profile.docx")
    except Exception as e:
        flash("Failed to generate DOCX. Please try again or contact support.")
        print(f"DOCX generation failed: {e}")
        return redirect('/')

@app.route('/download_xlsx', methods=['POST'])
def download_xlsx():
    profile = session.get('profile', {})
    
    # Get hidden sections from form data
    hidden_sections = []
    try:
        hidden_sections_str = request.form.get('hidden_sections', '[]')
        hidden_sections = json.loads(hidden_sections_str)
        if not isinstance(hidden_sections, list):
            print(f"Warning: hidden_sections is not a valid list, defaulting to empty list: {hidden_sections}")
            hidden_sections = []
        # Validate section IDs
        hidden_sections = [section for section in hidden_sections if section in VALID_SECTION_IDS]
    except json.JSONDecodeError as e:
        print(f"Error decoding hidden_sections: {e}, defaulting to []")
        hidden_sections = []
    
    print(f"Download XLSX request - hidden_sections: {hidden_sections}")
    
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /download_xlsx")
        return redirect('/')
    try:
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.xlsx")
        render_html_to_xlsx(profile, output_path, hidden_sections)
        print(f"Generated XLSX: {output_path} with hidden sections: {hidden_sections}")
        return send_file(output_path, as_attachment=True, download_name="Employee_Profile.xlsx")
    except Exception as e:
        flash("Failed to generate XLSX. Please try again or contact support.")
        print(f"XLSX generation failed: {e}")
        return redirect('/')

@app.route('/display_profile')
def display_profile():
    profile = session.get('profile', {})
    hidden_sections = session.get('hidden_sections', [])
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /display_profile")
        return redirect('/')
    print(f"Rendering display_profile.html with profile={json.dumps(profile, indent=2)}, hidden_sections={hidden_sections}")
    return render_template("display_profile.html", profile=profile, hidden_sections=hidden_sections)

# === Run ===
if __name__ == '__main__':
    app.run(debug=True)