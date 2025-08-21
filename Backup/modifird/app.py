import os
import uuid
import json
import re
from flask import Flask, render_template, request, send_file, flash, redirect, session
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import Workbook
import google.generativeai as genai
from playwright.sync_api import sync_playwright

# Gemini API Key
genai.configure(api_key="AIzaSyBWHfbDFcuhKwEL-uX0j-SRyUUSidgBhaE")

app = Flask(__name__)
app.secret_key = 'supersecret'
UPLOAD_FOLDER = 'Uploads'
GENERATED_FOLDER = 'generated'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

model = genai.GenerativeModel("gemini-2.5-flash")

# === Helper Functions ===
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(filepath, ext):
    try:
        if ext == 'pdf':
            with open(filepath, 'rb') as f:
                text = ' '.join([p.extract_text() for p in PdfReader(f).pages if p.extract_text()])
                print(f"Extracted PDF text: {text[:100]}...")
                return text
        elif ext == 'docx':
            text = '\n'.join([p.text for p in Document(filepath).paragraphs])
            print(f"Extracted DOCX text: {text[:100]}...")
            return text
        elif ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
                print(f"Extracted TXT text: {text[:100]}...")
                return text
        return ""
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""

def clean_formatting(text):
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'•', '-', text)
    text = re.sub(r'\n\s*-\s*', '\n- ', text)
    return text.strip()

def extract_json(text):
    try:
        if "```json" in text:
            match = re.findall(r"```json(.*?)```", text, re.DOTALL)
            if match:
                return json.loads(match[0].strip())
        return json.loads(text)
    except Exception as e:
        print(f"Error parsing JSON: {e}")
        return {}

def generate_structured_data(text):
    prompt = f"""
    You are an expert HR resume parser. Convert the following text into structured JSON with these fields:
    {{
      "name": "",
      "education_training_certifications": [],
      "total_experience": "",
      "professional_summary": "",
      "netweb_projects": [{{"title": "", "description": ""}}],
      "past_projects": [{{"title": "", "description": ""}}],
      "roles_responsibilities": "",
      "technical_skills": {{
        "web_technologies": [],
        "scripting_languages": [],
        "frameworks": [],
        "databases": [],
        "web_servers": [],
        "tools": []
      }},
      "personal_details": {{
        "employee_id": "",
        "permanent_address": "",
        "local_address": "",
        "contact_number": "",
        "date_of_joining": "",
        "designation": "",
        "overall_experience": "",
        "date_of_birth": "",
        "passport_details": ""
      }}
    }}
    Return ONLY valid JSON. Do not return markdown or explanation.
    Resume:
    {text}
    """
    try:
        response = model.generate_content(prompt)
        data = extract_json(response.text)
        print(f"Generated profile data: {json.dumps(data, indent=2)}")
        return data
    except Exception as e:
        print(f"Error generating structured data: {e}")
        return {}

def render_html_to_pdf(html_string, output_path):
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.set_content(html_string)
            page.pdf(path=output_path, format="A4", print_background=True)
            browser.close()
    except Exception as e:
        print(f"Error generating PDF: {e}")

def render_html_to_docx(profile, output_path):
    try:
        doc = Document()
        doc.add_heading('Professional Resume', 0)
        for key, value in profile.items():
            if isinstance(value, dict):
                doc.add_heading(key.replace('_', ' ').title(), level=1)
                for subkey, subvalue in value.items():
                    if subvalue:
                        doc.add_heading(subkey.replace('_', ' ').title(), level=2)
                        doc.add_paragraph(str(subvalue))
            elif isinstance(value, list):
                if value:
                    doc.add_heading(key.replace('_', ' ').title(), level=1)
                    for item in value:
                        doc.add_paragraph(str(item), style='ListBullet')
            else:
                if value:
                    doc.add_heading(key.replace('_', ' ').title(), level=1)
                    doc.add_paragraph(str(value))
        doc.save(output_path)
    except Exception as e:
        print(f"Error generating DOCX: {e}")

def render_html_to_xlsx(profile, output_path):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Resume"
        row = 1
        for key, value in profile.items():
            if value:
                ws.cell(row=row, column=1).value = key.replace('_', ' ').title()
                if isinstance(value, dict):
                    row += 1
                    for subkey, subvalue in value.items():
                        if subvalue:
                            ws.cell(row=row, column=2).value = subkey.replace('_', ' ').title()
                            ws.cell(row=row, column=3).value = str(subvalue)
                            row += 1
                elif isinstance(value, list):
                    row += 1
                    for item in value:
                        ws.cell(row=row, column=2).value = str(item)
                        row += 1
                else:
                    ws.cell(row=row, column=2).value = str(value)
                    row += 1
        wb.save(output_path)
    except Exception as e:
        print(f"Error generating XLSX: {e}")

# === Routes ===
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('file_input')
        prompt_text = request.form.get('text_input')
        text = ""
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            ext = filename.rsplit('.', 1)[1].lower()
            text = extract_text(path, ext)
        elif prompt_text and prompt_text.strip():
            text = prompt_text.strip()
        if not text:
            flash("Please upload a file or enter resume text.")
            return redirect('/')
        formatted_text = clean_formatting(text)
        profile = generate_structured_data(formatted_text)
        if not profile:
            flash("Failed to generate profile data.")
            return redirect('/')
        session['profile'] = profile
        session['hidden_sections'] = []
        print(f"Initialized session: profile={json.dumps(profile, indent=2)}, hidden_sections=[]")
        return render_template("profile.html", profile=profile, hidden_sections=[])
    session.pop('profile', None)
    session.pop('hidden_sections', None)
    return render_template("index.html")

@app.route('/update_profile', methods=['POST'])
def update_profile():
    action = request.form.get('action')
    print("Form data received:")
    form_data = request.form.to_dict(flat=False)
    for key, value in form_data.items():
        print(f"{key}: {value}")
    
    # Parse hidden sections
    hidden_sections = []
    try:
        hidden_sections = json.loads(request.form.get('hidden_sections', '[]'))
        if not isinstance(hidden_sections, list):
            print(f"Warning: hidden_sections is not a list, defaulting to []: {hidden_sections}")
            hidden_sections = []
    except json.JSONDecodeError as e:
        print(f"Error decoding hidden_sections: {e}, defaulting to []")
        hidden_sections = []

    profile_data = {
        "name": request.form.get('name', '').strip(),
        "education_training_certifications": [item.strip() for item in request.form.getlist('education_training_certifications[]') if item.strip()],
        "total_experience": request.form.get('total_experience', '').strip(),
        "professional_summary": request.form.get('professional_summary', '').strip(),
        "netweb_projects": [
            {"title": title.strip(), "description": desc.strip()}
            for title, desc in zip(
                request.form.getlist('netweb_projects[title][]'),
                request.form.getlist('netweb_projects[description][]')
            )
            if title.strip() or desc.strip()
        ],
        "past_projects": [
            {"title": title.strip(), "description": desc.strip()}
            for title, desc in zip(
                request.form.getlist('past_projects[title][]'),
                request.form.getlist('past_projects[description][]')
            )
            if title.strip() or desc.strip()
        ],
        "roles_responsibilities": request.form.get('roles_responsibilities', '').strip(),
        "technical_skills": {
            "web_technologies": [item.strip() for item in request.form.getlist('technical_skills[web_technologies][]') if item.strip()],
            "scripting_languages": [item.strip() for item in request.form.getlist('technical_skills[scripting_languages][]') if item.strip()],
            "frameworks": [item.strip() for item in request.form.getlist('technical_skills[frameworks][]') if item.strip()],
            "databases": [item.strip() for item in request.form.getlist('technical_skills[databases][]') if item.strip()],
            "web_servers": [item.strip() for item in request.form.getlist('technical_skills[web_servers][]') if item.strip()],
            "tools": [item.strip() for item in request.form.getlist('technical_skills[tools][]') if item.strip()]
        },
        "personal_details": {
            "employee_id": request.form.get('personal_details[employee_id]', '').strip(),
            "permanent_address": request.form.get('personal_details[permanent_address]', '').strip(),
            "local_address": request.form.get('personal_details[local_address]', '').strip(),
            "contact_number": request.form.get('personal_details[contact_number]', '').strip(),
            "date_of_joining": request.form.get('personal_details[date_of_joining]', '').strip(),
            "designation": request.form.get('personal_details[designation]', '').strip(),
            "overall_experience": request.form.get('personal_details[overall_experience]', '').strip(),
            "date_of_birth": request.form.get('personal_details[date_of_birth]', '').strip(),
            "passport_details": request.form.get('personal_details[passport_details]', '').strip()
        }
    }

    # Validate profile data
    has_data = (
        profile_data['name'] or
        profile_data['education_training_certifications'] or
        profile_data['total_experience'] or
        profile_data['professional_summary'] or
        profile_data['netweb_projects'] or
        profile_data['past_projects'] or
        profile_data['roles_responsibilities'] or
        any(profile_data['technical_skills'].values()) or
        any(profile_data['personal_details'].values())
    )
    if not has_data:
        flash("No valid data provided. Please fill in at least one field.")
        print("Validation failed: No valid data provided")
        return render_template("profile.html", profile=profile_data, hidden_sections=hidden_sections)

    session['profile'] = profile_data
    session['hidden_sections'] = hidden_sections
    print(f"Updated session profile: {json.dumps(profile_data, indent=2)}")
    print(f"Updated session hidden_sections: {hidden_sections}")

    if action == 'save':
        return render_template("display_profile.html", profile=profile_data, hidden_sections=hidden_sections)
    else:
        flash("Invalid action.")
        print("Invalid action received")
        return render_template("profile.html", profile=profile_data, hidden_sections=hidden_sections)

@app.route('/edit_profile')
def edit_profile():
    profile = session.get('profile', {})
    hidden_sections = session.get('hidden_sections', [])
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /edit_profile")
        return redirect('/')
    print(f"Edit profile: profile={json.dumps(profile, indent=2)}, hidden_sections={hidden_sections}")
    return render_template("profile.html", profile=profile, hidden_sections=hidden_sections)

@app.route('/download_pdf')
def download_pdf():
    profile = session.get('profile', {})
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /download_pdf")
        return redirect('/')
    try:
        html = render_template("display_profile.html", profile=profile, hidden_sections=[])
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.pdf")
        render_html_to_pdf(html, output_path)
        print(f"Generated PDF: {output_path}")
        return send_file(output_path, as_attachment=True, download_name="Employee_Profile.pdf")
    except Exception as e:
        flash(f"PDF generation failed: {str(e)}")
        print(f"PDF generation failed: {e}")
        return redirect('/')

@app.route('/download_docx')
def download_docx():
    profile = session.get('profile', {})
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /download_docx")
        return redirect('/')
    try:
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.docx")
        render_html_to_docx(profile, output_path)
        print(f"Generated DOCX: {output_path}")
        return send_file(output_path, as_attachment=True, download_name="Employee_Profile.docx")
    except Exception as e:
        flash(f"DOCX generation failed: {str(e)}")
        print(f"DOCX generation failed: {e}")
        return redirect('/')

@app.route('/download_xlsx')
def download_xlsx():
    profile = session.get('profile', {})
    if not profile:
        flash("No profile data available.")
        print("No profile data in session for /download_xlsx")
        return redirect('/')
    try:
        output_path = os.path.join(GENERATED_FOLDER, f"profile_{uuid.uuid4().hex}.xlsx")
        render_html_to_xlsx(profile, output_path)
        print(f"Generated XLSX: {output_path}")
        return send_file(output_path, as_attachment=True, download_name="Employee_Profile.xlsx")
    except Exception as e:
        flash(f"XLSX generation failed: {str(e)}")
        print(f"XLSX generation failed: {e}")
        return redirect('/')

# === Run ===
if __name__ == '__main__':
    app.run(debug=True)